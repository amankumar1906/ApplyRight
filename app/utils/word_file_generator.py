from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
import re
import os

def generate_word_resume(resume_text: str, output_path: str = "tailored_resume.docx") -> str:
    doc = Document()

    # — Global style & margins —
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(9.5)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Helpers
    SECTION_HEADERS = {
        "EDUCATION", "SKILLS", "TECHNICAL SKILLS",
        "PROFESSIONAL EXPERIENCE", "PROJECTS", "STARTUP INITIATIVES"
    }
    exp_count = 0
    proj_count = 0

    def split_date(text):
        m = re.search(r'(\d{4}.*)$', text)
        if m:
            return text[:m.start()].strip(), m.group(1).strip()
        return text, ""

    def flush_bullets(buffer):
        for txt in buffer:
            clean = txt.lstrip('●• ').strip()
            p = doc.add_paragraph(clean, style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
        buffer.clear()

    # Parse lines
    lines = [l.strip() for l in resume_text.splitlines() if l.strip()]
    header_block = []
    body_block = []
    seen_section = False
    for line in lines:
        if not seen_section and line.upper() in SECTION_HEADERS:
            seen_section = True
            body_block.append(line)
        elif not seen_section:
            header_block.append(line)
        else:
            body_block.append(line)

    # Render name/contact header
    for hl in header_block:
        p = doc.add_paragraph()
        run = p.add_run(hl)
        run.bold = True
        run.font.size = Pt(11)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Process body
    bullet_buf = []
    current_section = ""
    for line in body_block:
        uline = line.upper()

        # Section header
        if uline in SECTION_HEADERS:
            flush_bullets(bullet_buf)
            p = doc.add_paragraph()
            run = p.add_run(uline)
            run.bold = True
            run.font.size = Pt(10)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            current_section = uline
            continue

        # PROFESSIONAL EXPERIENCE entries
        if current_section == "PROFESSIONAL EXPERIENCE" and not line.startswith(('●','•')):
            flush_bullets(bullet_buf)
            exp_count += 1
            title, date = split_date(line)
            p = doc.add_paragraph()
            # set a right-aligned tab stop at ~6 inches
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
            run = p.add_run(f"{exp_count}. {title}")
            run.bold = True
            run.font.size = Pt(9.5)
            if date:
                run_tab = p.add_run('\t')
                date_run = p.add_run(date)
                date_run.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            continue

        # PROJECTS entries
        if current_section == "PROJECTS" and not line.startswith(('●','•')):
            flush_bullets(bullet_buf)
            proj_count += 1
            title, date = split_date(line)
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6), WD_TAB_ALIGNMENT.RIGHT)
            run = p.add_run(f"{proj_count}. {title}")
            run.bold = True
            run.font.size = Pt(9.5)
            if date:
                p.add_run('\t').font.size = Pt(9.5)
                p.add_run(date).font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            continue

        # SKILLS: bullet every line
        if current_section in ("SKILLS", "TECHNICAL SKILLS"):
            flush_bullets(bullet_buf)
            clean = line.lstrip('●• ').strip()
            p = doc.add_paragraph(clean, style='List Bullet')
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            continue

        # Bullet lines under Experience/Projects
        if line.startswith(('●','•')):
            bullet_buf.append(line)
            continue

        # Other sections (Education, Startup)
        flush_bullets(bullet_buf)
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

    flush_bullets(bullet_buf)
    doc.save(output_path)
    return os.path.abspath(output_path)
