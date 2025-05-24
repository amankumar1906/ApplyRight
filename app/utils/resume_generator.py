import textwrap
import asyncio
import logging
from contextlib import closing
import google.generativeai as genai
from app.core.config import Config
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Configure Gemini API
genai.configure(api_key=Config.GOOGLE_API_KEY)

def _generate_tailored_resume_sync(resume_text: str, jd_metadata: dict, reference_resumes: str = "") -> tuple[str, BytesIO]:
    """
    Synchronously generate a tailored resume and save it as a Word document in a BytesIO buffer.

    Args:
        resume_text (str): The original resume text.
        jd_metadata (dict): Job description metadata with skills and job title.
        reference_resumes (str, optional): Reference resumes for context.

    Returns:
        tuple[str, BytesIO]: Tailored resume text and the Word document buffer.

    Raises:
        ValueError: If input data is invalid or empty.
        Exception: If Gemini API or document creation fails.
    """
    # Validate inputs
    if not resume_text or not jd_metadata:
        logger.error("Resume text or JD metadata is empty")
        raise ValueError("Resume text or JD metadata is empty")

    # Format skills based on weight
    high = [f'{s["skill"]} ({s["weight"]})' for s in jd_metadata["skills"] if s["weight"] > 8]
    medium = [f'{s["skill"]} ({s["weight"]})' for s in jd_metadata["skills"] if 3 <= s["weight"] <= 8]
    low = [f'{s["skill"]} ({s["weight"]})' for s in jd_metadata["skills"] if s["weight"] < 3]
    def fmt(lst): return ", ".join(lst) if lst else "—"

    # Construct prompt
    SYSTEM = textwrap.dedent(f"""
        You are an elite, human-sounding résumé editor.

        ── OUTPUT SCOPE ─────────────
        • Return the entire résumé in this exact section order:

        Name and Address Details
        Education
        Skills (merge original + skills extracted from the job description): Group into exactly three bullet sub-headings using distinct categories (e.g., Front-end, Back-end, DevOps); list all skills in each group on a single line, comma-separated, without bullet points.
        Professional Experience
        Projects
        • Each section header appears once, followed by its content.
        • Include up to two projects. If creating a new project, remove the existing one least aligned with the job description.
        ── PRIORITY SKILLS ──────────
        • High emphasis (≥2 mentions): {fmt(high)}

        • Medium emphasis (≥1 mention): {fmt(medium)}

        • Low emphasis (nice-to-have): {fmt(low)}

        ── EXPERIENCE RULES ─────────

        • Retain all existing bullets but make sure they are 35-40 words; rephrase for clarity/metrics while preserving meaning.
        • If a role lacks a high/medium skill, add one new bullet (30-35 words).
        • All bullets must be human-sounding, metric-driven, and avoid AI filler.

        ── PROJECT RULES ────────────

        Evaluate existing projects against priority skills:
        • No priority skills → delete and create one new project (4 bullets).
        • Some priority skills → keep and refine bullets to include missing skills.
        • All priority skills → polish language and integrate skill keywords naturally.
        New projects must:
        • Be realistic (≤2 years old).
        • Include 3 bullets (18–25 words each).
        • Highlight as many priority skills/technologies as possible.
        Final résumé must have no more than two projects.
        ── STYLE & FORMAT ───────────
        • Use bullets starting with action verbs, followed by metrics/outcomes, in plain text (no Markdown/JSON).

        • Keep total length within ±15% of original, avoiding AI fluff phrases.

        • Self-audit: ensure every high skill appears ≥2×, every medium skill ≥1×.

        Respond ONLY with the final résumé, formatted per the above structure. Do not add commentary, markdown, or extra notes.
    """)
    USER = f"""--- ORIGINAL CONTENT START ---\n{resume_text}\n--- ORIGINAL CONTENT END ---"""
    if reference_resumes.strip():
        USER += f"""\n\n--- REFERENCE RESUMES WITH GUARANTEED HIGH SCORE ---\n{reference_resumes.strip()}\n--- END REFERENCES ---"""

    # Generate tailored resume text
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        reply = model.generate_content([SYSTEM, USER])
        tailored_resume_text = reply.text
        if not tailored_resume_text or tailored_resume_text.isspace():
            logger.error("Generated resume text is empty or invalid")
            raise ValueError("Generated resume text is empty or invalid")
    except Exception as e:
        logger.error(f"Failed to generate tailored resume: {e}")
        raise

    # Create Word document with professional styling
    buffer = BytesIO()
    try:
        doc = Document()

        # Set document margins to 0.5 inches
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Define styles
        styles = doc.styles
        if 'Normal' in styles:
            normal_style = styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.space_after = Pt(4)
            normal_style.paragraph_format.line_spacing = 1.0

        # Define custom style for section headers
        if 'Heading 2' not in styles:
            styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        header_style = styles['Heading 2']
        header_style.font.name = 'Times New Roman'
        header_style.font.size = Pt(14)
        header_style.font.bold = True
        header_style.paragraph_format.space_before = Pt(8)
        header_style.paragraph_format.space_after = Pt(4)

        # Define custom style for bullet points
        bullet_style = styles.add_style('ResumeBullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Times New Roman'
        bullet_style.font.size = Pt(11)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
        bullet_style.paragraph_format.first_line_indent = Inches(-0.25)
        bullet_style.paragraph_format.space_after = Pt(3)
        bullet_style.paragraph_format.line_spacing = 1.0

        # Process the resume text
        lines = tailored_resume_text.strip().splitlines()
        current_section = None
        in_role_or_project = False
        role_or_project_lines = []

        # Add name and contact details
        if lines:
            name = lines[0].strip()
            contact_lines = []
            i = 1
            while i < len(lines) and not any(lines[i].lower().startswith(section.lower()) for section in ["education", "skills", "professional experience", "projects"]):
                contact_lines.append(lines[i].strip())
                i += 1

            # Add name (centered, bold, 16pt)
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(name)
            name_run.font.name = 'Times New Roman'
            name_run.font.size = Pt(16)
            name_run.bold = True

            # Add contact details (centered, 10pt)
            if contact_lines:
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_text = " | ".join(contact_lines)
                contact_run = contact_para.add_run(contact_text)
                contact_run.font.name = 'Times New Roman'
                contact_run.font.size = Pt(10)
                contact_para.paragraph_format.space_after = Pt(8)

        # Process remaining sections
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # Detect section headers
            if any(line.lower().startswith(section.lower()) for section in ["education", "skills", "professional experience", "projects"]):
                current_section = line
                doc.add_paragraph(current_section, style='Heading 2')
                in_role_or_project = False
                role_or_project_lines = []
                i += 1
                continue

            # Handle role/project titles (e.g., "Software Engineer Intern: READY.NET, Boston, MA May - August 2023")
            if current_section and current_section.lower() == "professional experience" or current_section.lower() == "projects":
                if not line.startswith("- ") and not in_role_or_project:
                    role_or_project_lines.append(line)
                    if "May" in line or "January" in line or "June" in line:  # Assume dates indicate end of title block
                        in_role_or_project = True
                        title = " | ".join(role_or_project_lines)
                        para = doc.add_paragraph(title)
                        para.style.font.name = 'Times New Roman'
                        para.style.font.size = Pt(11)
                        para.paragraph_format.space_after = Pt(3)
                        role_or_project_lines = []
                    i += 1
                    continue

            # Handle bullets and other content
            if line.startswith("- "):
                para = doc.add_paragraph(line[2:], style='ResumeBullet')
            else:
                para = doc.add_paragraph(line)
                para.style.font.name = 'Times New Roman'
                para.style.font.size = Pt(11)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.line_spacing = 1.0

            i += 1

        doc.save(buffer)
        buffer.seek(0)
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        buffer.close()
        raise
    return tailored_resume_text, buffer

async def generate_tailored_resume(resume_text: str, jd_metadata: dict, reference_resumes: str = "") -> tuple[str, BytesIO]:
    """
    Asynchronously generate a tailored resume by running the sync function in a thread.

    Args:
        resume_text (str): The original resume text.
        jd_metadata (dict): Job description metadata.
        reference_resumes (str, optional): Reference resumes for context.

    Returns:
        tuple[str, BytesIO]: Tailored resume text and the Word document buffer.
    """
    return await asyncio.to_thread(
        _generate_tailored_resume_sync,
        resume_text,
        jd_metadata,
        reference_resumes
    )